# la idea es automatizar algunos registros que estamos realizando en Hopex, nuestra herramienta.
# quien esta a cargo de mi (Yeyson) requiere actualizar algunos procesos, son 4000, pero en 1 día solo he logrado hacer 168
# mi idea (Carlos Peña M) es automatizar dicho proceso con algun programa de python
import os
import time
from selenium import webdriver
from selenium.webdriver.common.keys import keys
import datetime
import pyautogui
from docx import Document
from docx.shared import Inches
from openpxyl import load_workbook

#Con esto vamos a configurar Selenium y Microsoft Edge
path="C:\Users\carpena\Downloads"
driver=webdriver.Edge(executable_path=path)
driver.maximize_window()

# cargar el archivo de excel, esté podrá ser un archivo de control que exportemos desde hopex
excel_file="miarchivo.xlsx"
workbook=load_workbook(excel_file)
worksheet=workbook.active
event_list=[cel.value for cell in worksheet["A"]]


#configuración del documento de word
# podriamos utilizar pyautogui.screenshot() para capturar los movimientos realizados y auditar que todos los registros y procesos se hayan realizado automaticamente

fecha_actual=datetime.datetime.now()
nombre_archivo=fecha_actual.strftime("%d-%m-%Y")+' - Carlos A Peña Montenegro -'+'reporte.docx'
if os.path.exists(nombre_archivo):
    doc=Document(nombre_archivo)
else:
    doc=Document()
    doc.save(nombre_archivo)

document=Document()
document.add_heading("Reporte"+time.strftime("%d-%m-%Y"),0)


pyautogui.moveTo(x=100, y=100)
pyautogui.rightClick()

driver=webdriver.Edge()
driver.get('https://www.bancolombia.com/')



doc=Document()

